from copy import deepcopy
from pathlib import Path
from tempfile import NamedTemporaryFile
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt


REFERENCE = Path("/Users/minifu/Project/AI_Gender/.codex_tmp/cv_yuhao/template.docx")
OUTPUT = Path("/Users/minifu/Project/AI_Gender/Yuhao_Fu_CV.docx")


def set_run(run, *, bold=None, italic=None, underline=None):
    run.font.name = "Century"
    run.font.size = Pt(10.5)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if underline is not None:
        run.underline = underline
    return run


def add_para(doc, text="", style="Normal", *, bold=False, first_indent=None,
             left_indent=None, space_before=0, space_after=0, keep_next=False):
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if first_indent is not None:
        p.paragraph_format.first_line_indent = Inches(first_indent)
    if left_indent is not None:
        p.paragraph_format.left_indent = Inches(left_indent)
    p.paragraph_format.keep_with_next = keep_next
    if text:
        set_run(p.add_run(text), bold=bold)
    return p


doc = Document(REFERENCE)
body = doc._element.body
sect_pr = body.sectPr
for child in list(body):
    if child is not sect_pr:
        body.remove(child)

# Identity block
add_para(doc, "Yuhao Fu", "Heading", bold=True, keep_next=True)
add_para(doc, "PhD Student", keep_next=True)
add_para(doc, "Graduate School of Economics", keep_next=True)
add_para(doc, "The University of Osaka")
add_para(doc)

# Current position
add_para(doc, "Current Position", "Heading 5", bold=True, keep_next=True)
add_para(doc, "April 2024 – Present", keep_next=True)
add_para(
    doc,
    "PhD Student, Graduate School of Economics, The University of Osaka",
    first_indent=0.5,
)
add_para(doc)

# Education (highest completed degree only)
add_para(doc, "Education", "Heading 5", bold=True, keep_next=True)
p = add_para(doc)
p.paragraph_format.keep_with_next = False
set_run(p.add_run("2024\t"))
set_run(p.add_run("Master of Economics, The University of Osaka"))
tabs = p.paragraph_format.tab_stops
tabs.add_tab_stop(Inches(0.65))
add_para(doc)

# Sole journal publication
add_para(doc, "Journal Publication", "Heading 5", bold=True, keep_next=True)
p = add_para(doc, left_indent=0.30, first_indent=-0.30)
set_run(p.add_run("Fu, Y., and N. Hanaki. "))
set_run(p.add_run("“Do people rely on ChatGPT more than their peers to detect deepfake news?” "))
set_run(p.add_run("Journal of Economic Interaction and Coordination"), italic=True)
set_run(p.add_run(", 2026. https://doi.org/10.1007/s11403-026-00490-6"))

# Competitive doctoral fellowship and research funding
add_para(doc)
add_para(doc, "Honors, Fellowships, and Research Funding", "Heading 5", bold=True, keep_next=True)
p = add_para(doc)
p.paragraph_format.keep_with_next = False
set_run(p.add_run("2024–Present\t"))
set_run(
    p.add_run(
        "Recipient, JST Support for Pioneering Research Initiated by the Next "
        "Generation (SPRING), The University of Osaka"
    )
)
p.paragraph_format.tab_stops.add_tab_stop(Inches(0.95))

# Ensure section properties remain the last child in the body.
if body.sectPr is not None and body[-1] is not body.sectPr:
    body.remove(body.sectPr)
    body.append(sect_pr)

with NamedTemporaryFile(suffix=".docx", delete=False) as interim_file:
    interim = Path(interim_file.name)
doc.save(interim)

# Repackage from the retained reference so unrelated OOXML parts remain
# byte-for-byte source-derived. Only the body document part is replaced.
with ZipFile(REFERENCE, "r") as source, ZipFile(OUTPUT, "w", ZIP_DEFLATED) as target:
    with ZipFile(interim, "r") as generated:
        generated_document = generated.read("word/document.xml")
    for item in source.infolist():
        data = generated_document if item.filename == "word/document.xml" else source.read(item.filename)
        target.writestr(item, data)
interim.unlink()
print(OUTPUT)
